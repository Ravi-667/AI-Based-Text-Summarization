"""
Export utilities for generating PDF and DOCX summaries
"""

from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional
import io

# PDF generation
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except ImportError:
    REPORTLAB_AVAILABLE = False

# DOCX generation
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class ExportManager:
    """Manager for exporting summaries to various formats."""
    
    @staticmethod
    def export_to_pdf(
        summary: str,
        original_text: str,
        metadata: Dict[str, Any],
        output_path: str
    ) -> bool:
        """
        Export summary to PDF format.
        
        Args:
            summary: Generated summary text
            original_text: Original text
            metadata: Summary metadata (method, stats, etc.)
            output_path: Output file path
            
        Returns:
            True if successful, False otherwise
        """
        if not REPORTLAB_AVAILABLE:
            raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
        
        try:
            # Create PDF document
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1f77b4'),
                spaceAfter=30,
                alignment=1  # Center
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#2ca02c'),
                spaceAfter=12
            )
            
            # Title
            title = Paragraph("AI-Generated Text Summary", title_style)
            story.append(title)
            story.append(Spacer(1, 0.3 * inch))
            
            # Metadata table
            metadata_data = [
                ['Generation Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')],
                ['Method', metadata.get('method', 'N/A').capitalize()],
                ['Original Words', str(metadata.get('original_length', 'N/A'))],
                ['Summary Words', str(metadata.get('summary_length', 'N/A'))],
                ['Compression', f"{metadata.get('compression_ratio', 0):.1f}%"]
            ]
            
            metadata_table = Table(metadata_data, colWidths=[2*inch, 3*inch])
            metadata_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
                ('TOPPADDING', (0, 0), (-1, -1), 8),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey)
            ]))
            
            story.append(metadata_table)
            story.append(Spacer(1, 0.4 * inch))
            
            # Summary section
            summary_heading = Paragraph("Summary", heading_style)
            story.append(summary_heading)
            
            summary_para = Paragraph(summary, styles['BodyText'])
            story.append(summary_para)
            story.append(Spacer(1, 0.3 * inch))
            
            # Original text section (truncated if too long)
            original_heading = Paragraph("Original Text", heading_style)
            story.append(original_heading)
            
            original_display = original_text[:2000] + "..." if len(original_text) > 2000 else original_text
            original_para = Paragraph(original_display, styles['BodyText'])
            story.append(original_para)
            
            # Build PDF
            doc.build(story)
            return True
            
        except Exception as e:
            print(f"Error exporting to PDF: {e}")
            return False
    
    @staticmethod
    def export_to_docx(
        summary: str,
        original_text: str,
        metadata: Dict[str, Any],
        output_path: str
    ) -> bool:
        """
        Export summary to DOCX format.
        
        Args:
            summary: Generated summary text
            original_text: Original text
            metadata: Summary metadata
            output_path: Output file path
            
        Returns:
            True if successful, False otherwise
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
        
        try:
            # Create document
            doc = Document()
            
            # Title
            title = doc.add_heading('AI-Generated Text Summary', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Metadata section
            doc.add_heading('Summary Information', level=2)
            
            metadata_table = doc.add_table(rows=5, cols=2)
            metadata_table.style = 'Light Grid Accent 1'
            
            metadata_items = [
                ('Generation Date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                ('Method', metadata.get('method', 'N/A').capitalize()),
                ('Original Words', str(metadata.get('original_length', 'N/A'))),
                ('Summary Words', str(metadata.get('summary_length', 'N/A'))),
                ('Compression', f"{metadata.get('compression_ratio', 0):.1f}%")
            ]
            
            for i, (key, value) in enumerate(metadata_items):
                row_cells = metadata_table.rows[i].cells
                row_cells[0].text = key
                row_cells[1].text = value
                # Make key cells bold
                for paragraph in row_cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            doc.add_paragraph()  # Spacer
            
            # Summary section
            doc.add_heading('Summary', level=2)
            summary_para = doc.add_paragraph(summary)
            summary_para.style = 'Body Text'
            
            doc.add_paragraph()  # Spacer
            
            # Original text section (truncated)
            doc.add_heading('Original Text', level=2)
            original_display = original_text[:2000] + "..." if len(original_text) > 2000 else original_text
            original_para = doc.add_paragraph(original_display)
            original_para.style = 'Body Text'
            
            # Footer
            doc.add_paragraph()
            footer = doc.add_paragraph('Generated by AI-Based Text Summarization System')
            footer.style = 'Subtle Emphasis'
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Save document
            doc.save(output_path)
            return True
            
        except Exception as e:
            print(f"Error exporting to DOCX: {e}")
            return False
    
    @staticmethod
    def export_to_txt(
        summary: str,
        original_text: str,
        metadata: Dict[str, Any],
        output_path: str
    ) -> bool:
        """
        Export summary to plain text format.
        
        Args:
            summary: Generated summary
            original_text: Original text
            metadata: Summary metadata
            output_path: Output file path
            
        Returns:
            True if successful
        """
        try:
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("=" * 80 + "\n")
                f.write("AI-GENERATED TEXT SUMMARY\n")
                f.write("=" * 80 + "\n\n")
                
                f.write("SUMMARY INFORMATION\n")
                f.write("-" * 40 + "\n")
                f.write(f"Generation Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                f.write(f"Method: {metadata.get('method', 'N/A').capitalize()}\n")
                f.write(f"Original Words: {metadata.get('original_length', 'N/A')}\n")
                f.write(f"Summary Words: {metadata.get('summary_length', 'N/A')}\n")
                f.write(f"Compression: {metadata.get('compression_ratio', 0):.1f}%\n\n")
                
                f.write("=" * 80 + "\n")
                f.write("SUMMARY\n")
                f.write("=" * 80 + "\n\n")
                f.write(summary + "\n\n")
                
                f.write("=" * 80 + "\n")
                f.write("ORIGINAL TEXT\n")
                f.write("=" * 80 + "\n\n")
                f.write(original_text + "\n")
            
            return True
            
        except Exception as e:
            print(f"Error exporting to TXT: {e}")
            return False
